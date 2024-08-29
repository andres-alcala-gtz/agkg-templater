import re
import docx
import pandas
import pathlib
import datetime
import geopy.geocoders


def search_part_style(path: pathlib.Path) -> dict[str, str]:

    part_paragraph = {"section": "Plan Appraisal - Sample Scope: Timeline and Scheduling", "title": "Appraisal Timeline", "subtitle": "Phase 2: Conduct Appraisal", "label": "Appraisal Type", "text": "Benchmark"}

    style_paragraphs = {}
    document = docx.Document(str(path))
    for paragraph in document.paragraphs:
        if paragraph.style is not None and paragraph.text != "":
            if paragraph.style.name not in style_paragraphs:
                style_paragraphs[paragraph.style.name] = [paragraph.text]
            else:
                style_paragraphs[paragraph.style.name].append(paragraph.text)

    part_style = {}
    for part, paragraph in part_paragraph.items():
        for style, paragraphs in style_paragraphs.items():
            if paragraph in paragraphs:
                part_style[part] = style

    return part_style


def search_data_frame(path: pathlib.Path, part_style: dict[str, str]) -> pandas.DataFrame:

    def _push(data_frame: pandas.DataFrame, dictionary: dict[str, str], mode: str) -> None:
        if mode == "a":
            data_frame.loc[len(data_frame)] = list(dictionary.values())
        if mode == "w":
            data_frame.loc[len(data_frame) - 1, part_style["text"]] += f'\n{dictionary[part_style["text"]]}'

    dictionary = {part_style["section"]: "", part_style["title"]: "", part_style["subtitle"]: "", part_style["label"]: "", part_style["text"]: ""}
    data_frame = pandas.DataFrame(columns=list(dictionary.keys()))

    dictionary.update({part_style["section"]: "Plan Appraisal - Cover", part_style["label"]: "AID:", part_style["text"]: re.search("[\d]+", path.name).group()})
    _push(data_frame, dictionary, "a")

    document = docx.Document(str(path))
    for paragraph in document.paragraphs:
        if paragraph.style is not None and paragraph.text != "":
            if paragraph.style.name in dictionary:
                dictionary.update({paragraph.style.name: paragraph.text})
            if paragraph.style.name == part_style["text"]:
                mode = "a" if list(data_frame.loc[len(data_frame) - 1])[:-1] != list(dictionary.values())[:-1] else "w"
                _push(data_frame, dictionary, mode)

    return data_frame


def search_data(data_frame: pandas.DataFrame, part_style: dict[str, str]) -> dict[str, str]:

    def _loc(personnel: bool = False, index: int = 0, section: str = None, title: str = None, subtitle: str = None, label: str = None, text: str = None) -> str:
        filtered = data_frame
        if section: filtered = filtered[filtered[part_style["section"]] == section]
        if title: filtered = filtered[filtered[part_style["title"]] == title]
        if subtitle: filtered = filtered[filtered[part_style["subtitle"]] == subtitle]
        if label: filtered = filtered[filtered[part_style["label"]] == label]
        if text: filtered = filtered[filtered[part_style["text"]] == text]
        try: return filtered[part_style["subtitle" if personnel else "text"]].values[index]
        except: return "-"

    data = {}

    data["«identifier»"] = _loc(section="Plan Appraisal - Cover", label="AID:")
    data["«target»"] = _loc(section="Plan Appraisal - Cover", label="Target:").replace("- ", "")
    data["«TARGET»"] = data["«target»"].upper()

    data["«partner»"] = _loc(section="Plan Appraisal – Sample Scope: Appraisal Setup", label="Partner")
    data["«PARTNER»"] = data["«partner»"].upper()
    data["«objectives»"] = _loc(section="Plan Appraisal – Sample Scope: Appraisal Setup", label="Business and Appraisal Objectives")
    _, data["«objectives_business»"], data["«objectives_appraisal»"] = re.split("\w+ objectives:\n", data["«objectives»"])

    data["«organization_name»"] = _loc(section="Plan Appraisal – Sample Scope: Organization", title="Organization", label="Name")
    data["«ORGANIZATION_NAME»"] = data["«organization_name»"].upper()
    data["«organization_name_native»"] = _loc(section="Plan Appraisal – Sample Scope: Organization", title="Organization", label="Native Language Name")
    data["«ORGANIZATION_NAME_NATIVE»"] = data["«organization_name_native»"].upper()
    data["«organization_city»"] = _loc(section="Plan Appraisal – Sample Scope: Organization", title="Organization", label="City")
    data["«organization_region»"] = _loc(section="Plan Appraisal – Sample Scope: Organization", title="Organization", label="State/Province/Region")
    data["«organization_country»"] = geopy.geocoders.Nominatim(user_agent="templater").geocode(query=data["«organization_city»"], language="en").address.split(", ")[-1]

    data["«ou_name»"] = _loc(section="Plan Appraisal – Sample Scope: Organization", title="Organizational Unit (OU)", label="Name")
    data["«OU_NAME»"] = data["«ou_name»"].upper()
    data["«ou_name_native»"] = _loc(section="Plan Appraisal – Sample Scope: Organization", title="Organizational Unit (OU)", label="Native Language Name")
    data["«OU_NAME_NATIVE»"] = data["«ou_name_native»"].upper()

    data["«sponsor_name»"] = _loc(personnel=True, section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", text="Appraisal Sponsor")
    data["«Sponsor_Name»"] = data["«sponsor_name»"].title()
    data["«sponsor_organization»"] = _loc(section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", subtitle=data["«sponsor_name»"], label="Organization")

    data["«ri_name»"] = _loc(personnel=True, section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", text="Registered Interpreter")
    data["«RI_Name»"] = data["«ri_name»"].title()
    data["«ri_organization»"] = _loc(section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", subtitle=data["«ri_name»"], label="Organization")

    data["«ouc_name»"] = _loc(personnel=True, section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", text="OU Coordinator")
    data["«OUC_Name»"] = data["«ouc_name»"].title()
    data["«ouc_organization»"] = _loc(section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", subtitle=data["«ouc_name»"], label="Organization")

    data["«atm0_name»"] = _loc(personnel=True, index=0, section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", text="Appraisal Team Member")
    data["«ATM0_Name»"] = data["«atm0_name»"].title()
    data["«atm0_organization»"] = _loc(section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", subtitle=data["«atm0_name»"], label="Organization")

    data["«atm1_name»"] = _loc(personnel=True, index=1, section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", text="Appraisal Team Member")
    data["«ATM1_Name»"] = data["«atm1_name»"].title()
    data["«atm1_organization»"] = _loc(section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", subtitle=data["«atm1_name»"], label="Organization")

    data["«atm2_name»"] = _loc(personnel=True, index=2, section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", text="Appraisal Team Member")
    data["«ATM2_Name»"] = data["«atm2_name»"].title()
    data["«atm2_organization»"] = _loc(section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", subtitle=data["«atm2_name»"], label="Organization")

    data["«atm3_name»"] = _loc(personnel=True, index=3, section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", text="Appraisal Team Member")
    data["«ATM3_Name»"] = data["«atm3_name»"].title()
    data["«atm3_organization»"] = _loc(section="Plan Appraisal - Sample Scope: Appraisal Personnel", title="Create Appraisal Team", subtitle=data["«atm3_name»"], label="Organization")

    data["«pinyin»"] = "(pinyin)" if data["«organization_country»"] == "China" else ""
    data["«Pinyin»"] = data["«pinyin»"].title()
    data["«PINYIN»"] = data["«pinyin»"].upper()

    plan_start = datetime.datetime.strptime(_loc(section="Plan Appraisal - Sample Scope: Timeline and Scheduling", label="Plan Appraisal Start Date"), "%Y/%m/%d")
    plan_end = datetime.datetime.strptime(_loc(section="Plan Appraisal - Sample Scope: Timeline and Scheduling", label="Plan Appraisal End Date"), "%Y/%m/%d")
    conduct_start = datetime.datetime.strptime(_loc(section="Plan Appraisal - Sample Scope: Timeline and Scheduling", label="Conduct Appraisal Start Date"), "%Y/%m/%d")
    conduct_end = datetime.datetime.strptime(_loc(section="Plan Appraisal - Sample Scope: Timeline and Scheduling", label="Conduct Appraisal End Date"), "%Y/%m/%d")
    expiration = conduct_end.replace(year=conduct_end.year + 3)

    data["«plan_start_date»"] = plan_start.strftime("%d-%b-%Y")
    data["«plan_end_date»"] = plan_end.strftime("%d-%b-%Y")
    data["«conduct_start_date»"] = conduct_start.strftime("%d-%b-%Y")
    data["«conduct_end_date»"] = conduct_end.strftime("%d-%b-%Y")
    data["«expiration_date»"] = expiration.strftime("%d-%b-%Y")

    data["«conduct_start_year»"] = conduct_start.strftime("%Y")
    data["«conduct_start_month»"] = conduct_start.strftime("%m")
    data["«conduct_start_day»"] = conduct_start.strftime("%d")

    data["«conduct_end_year»"] = conduct_end.strftime("%Y")
    data["«conduct_end_month»"] = conduct_end.strftime("%m")
    data["«conduct_end_day»"] = conduct_end.strftime("%d")

    domains = {"Development": "DEV", "Services": "SVC", "Suppliers": "SPM", "Virtual": "VRT", "Security": "SEC"}
    parts = data["«target»"].split()

    data["«target_short»"] = f"{parts[0]}-{domains[parts[1]]} {parts[2][0]}{parts[3][0]}{parts[4]}"
    data["«TARGET_SHORT»"] = data["«target_short»"].upper()

    return data


def search(directory_src: pathlib.Path) -> dict[str, str]:

    path = [path for path in directory_src.glob("*") if path.suffix == ".docx" and not path.name.startswith(".")][0]

    part_style = search_part_style(path)
    data_frame = search_data_frame(path, part_style)
    data = search_data(data_frame, part_style)

    return data
